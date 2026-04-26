"""
AI Contract Review — Fixed UI
"""
import os
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"

import sys, json, datetime, sqlite3
import streamlit as st

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from dotenv import load_dotenv
load_dotenv()

st.set_page_config(
    page_title="AI Contract Review",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap');

/* --- 1. CORE DISGUISE (HIDE STREAMLIT BRANDING) --- */
#MainMenu {visibility: hidden;}
header {visibility: hidden;}
footer {visibility: hidden;}

/* Base typography to custom modern font */
html, body, [class*="css"] {
    font-family: 'Outfit', -apple-system, sans-serif !important;
}

/* --- PRINT MODE CSS (CLEAN PDF EXPORTS) --- */
@media print {
    section[data-testid="stSidebar"], 
    header, 
    footer,
    iframe {
        display: none !important;
    }
    .stApp {
        background: #0f172a !important;
    }
    .block-container {
        max-width: 100% !important;
        padding-top: 10px !important;
        padding-left: 10px !important;
        padding-right: 10px !important;
    }
    .st-emotion-cache-1jicfl2, .st-emotion-cache-1104q3m { 
        width: 100% !important;
    }
}

/* --- 2. GLOBAL BACKGROUND (THE SAAS LOOK) --- */
.stApp {
    background: radial-gradient(125% 125% at 50% 0%, #0a0a0a 40%, #172554 100%) !important;
    color: #f3f4f6 !important;
}

/* --- 3. PREMIUM SIDEBAR --- */
section[data-testid="stSidebar"] {
    background: rgba(10, 15, 30, 0.6) !important;
    backdrop-filter: blur(20px) !important;
    border-right: 1px solid rgba(255, 255, 255, 0.05) !important;
}

/* Sidebar navigation buttons */
div[data-testid="stSidebar"] .stButton > button {
    width: 100%; 
    text-align: left; 
    background: transparent;
    border: none; 
    padding: 14px 20px; 
    border-radius: 12px;
    font-size: 15px; 
    font-weight: 500; 
    margin-bottom: 8px;
    color: #9ca3af;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
}
div[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(59, 130, 246, 0.1); 
    color: #60a5fa !important;
    transform: translateX(6px);
    box-shadow: -4px 0px 0px 0px #3b82f6;
}
/* Active Sidebar Button */
div[data-testid="stSidebar"] .stButton > button[kind="primary"] {
    background: linear-gradient(90deg, rgba(59, 130, 246, 0.2) 0%, transparent 100%);
    color: #ffffff !important;
    box-shadow: -4px 0px 0px 0px #60a5fa;
}

/* --- 4. GLASSMORPHISM METRIC CARDS --- */
[data-testid="stMetricValue"] {
    font-size: 32px !important;
    font-weight: 700 !important;
    background: linear-gradient(180deg, #ffffff 0%, #a1a1aa 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}
[data-testid="stMetricLabel"] {
    font-size: 13px !important;
    text-transform: uppercase;
    letter-spacing: 1px;
    color: #9ca3af !important;
}
[data-testid="stMetric"] {
    background: linear-gradient(135deg, rgba(17, 24, 39, 0.8) 0%, rgba(10, 15, 30, 0.8) 100%);
    backdrop-filter: blur(16px);
    border: 1px solid rgba(255, 255, 255, 0.08);
    border-radius: 20px;
    padding: 30px;
    box-shadow: 0 10px 30px -5px rgba(0, 0, 0, 0.6);
    transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
}
[data-testid="stMetric"]:hover {
    transform: translateY(-6px) scale(1.02);
    box-shadow: 0 20px 40px -10px rgba(59, 130, 246, 0.3);
    border: 1px solid rgba(59, 130, 246, 0.4);
}
.dashboard-card {
    background: rgba(15, 23, 42, 0.4);
    border: 1px solid rgba(255, 255, 255, 0.05);
    border-radius: 24px;
    padding: 32px;
    height: 100%;
}

/* --- 5. NEON BUTTONS --- */
div.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #2563eb 0%, #4f46e5 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 12px 24px !important;
    font-weight: 600 !important;
    box-shadow: 0 4px 12px rgba(37, 99, 235, 0.3) !important;
    transition: all 0.3s ease !important;
}
div.stButton > button[kind="primary"]:hover {
    transform: translateY(-2px) scale(1.02) !important;
    box-shadow: 0 8px 24px rgba(79, 70, 229, 0.5) !important;
}

/* --- 6. MODERN INPUT FIELDS --- */
.stTextArea textarea, .stTextInput input {
    border-radius: 12px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    background: rgba(15, 23, 42, 0.6) !important;
    color: #ffffff !important;
    padding: 16px !important;
    font-size: 15px !important;
    transition: all 0.2s ease !important;
}
.stTextArea textarea:focus, .stTextInput input:focus {
    border-color: #60a5fa !important;
    background: rgba(15, 23, 42, 0.9) !important;
    box-shadow: 0 0 0 4px rgba(59, 130, 246, 0.15) !important;
}

/* --- 7. MASTERS-WORTHY CHAT INTERFACE --- */
.chat-container {
    padding: 20px 0;
    max-width: 900px;
    margin: 0 auto;
}
.user-msg {
    background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-radius: 20px 20px 4px 20px;
    padding: 16px 24px;
    margin-bottom: 20px;
    box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    position: relative;
}
.bot-msg {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.4) 0%, rgba(15, 23, 42, 0.4) 100%);
    backdrop-filter: blur(12px);
    border: 1px solid rgba(96, 165, 250, 0.2);
    border-left: 5px solid #3b82f6;
    border-radius: 20px 20px 20px 4px;
    padding: 20px 28px;
    margin-bottom: 30px;
    box-shadow: 0 8px 32px rgba(0,0,0,0.3);
}
.msg-label {
    font-size: 11px;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    color: #60a5fa;
    margin-bottom: 8px;
    font-weight: 700;
}
.source-tag {
    background: rgba(59, 130, 246, 0.1);
    color: #93c5fd;
    padding: 2px 8px;
    border-radius: 4px;
    font-size: 11px;
    margin-right: 6px;
    border: 1px solid rgba(59, 130, 246, 0.3);
}
.glass-card {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.4) 0%, rgba(15, 23, 42, 0.4) 100%);
    backdrop-filter: blur(20px);
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-radius: 24px;
    padding: 35px;
    box-shadow: 0 15px 35px rgba(0,0,0,0.4);
    margin-bottom: 25px;
}
.executive-header {
    background: linear-gradient(90deg, #3b82f6 0%, #6366f1 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-weight: 800 !important;
    letter-spacing: -1px;
}
.metric-value-large {
    font-size: 3rem;
    font-weight: 500;
    color: #f8fafc;
    line-height: 1;
    letter-spacing: -1px;
}
.metric-label-small {
    font-size: 0.75rem;
    color: #94a3b8;
    text-transform: uppercase;
    letter-spacing: 2px;
}

.audit-row {
    padding: 12px 16px; 
    border-radius: 12px;
    background: rgba(16, 185, 129, 0.05); 
    border: 1px solid rgba(16, 185, 129, 0.1);
    margin-bottom: 8px; 
    font-size: 14px;
    color: #6ee7b7;
    transition: all 0.2s;
}
.audit-row:hover {
    background: rgba(16, 185, 129, 0.1);
    transform: translateX(4px);
}

/* Clean up expanders */
.streamlit-expanderHeader {
    font-size: 16px !important;
    font-weight: 500 !important;
    color: #e5e7eb !important;
    border-radius: 12px !important;
}
div[data-testid="stExpander"] {
    background: rgba(255,255,255,0.02) !important;
    border: 1px solid rgba(255,255,255,0.05) !important;
    border-radius: 12px !important;
}

/* Typography Overrides */
h1, h2, h3 {
    color: #ffffff !important;
    font-weight: 600 !important;
    letter-spacing: -0.03em;
}
/* --- 8. SMOOTH PAGE TRANSITIONS --- */
@keyframes fadeIn {
    from { opacity: 0; transform: translateY(10px); }
    to { opacity: 1; transform: translateY(0); }
}
.stApp {
    animation: fadeIn 0.5s ease-out;
}

hr {
    border-color: rgba(255, 255, 255, 0.05) !important;
}

/* --- 9. FIX COMPONENT IFRAME BACKGROUNDS --- */
iframe {
    background-color: transparent !important;
    border: none !important;
}
</style>
""", unsafe_allow_html=True)

# ── Session state ─────────────────────────────────────────────────────────────
_defaults = {
    "page": "upload", "contract_name": "", "chunks": [], "match_results": [],
    "topic_results": [], "contract_analysis": None, "indexed": False,
    "chat_history": [], "hitl_decisions": {}, "audit_log": [],
    "prefill_q": "", "auto_ask": False, "dynamic_questions": [],
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 📋 AI Contract Review")
    st.caption("AI-powered contract analysis")
    st.divider()
    if st.session_state.contract_name:
        st.success(f"📄 {st.session_state.contract_name[:28]}")
    else:
        st.info("No contract loaded")
    st.divider()
    for pid, label in [
        ("upload",    "📤  Upload Contract"),
        ("review",    "🔍  Review Contract"),
        ("qa",        "💬  Q&A"),
        ("dashboard", "📊  Dashboard"),
    ]:
        if st.button(label, key=f"nav_{pid}",
                     type="primary" if st.session_state.page == pid else "secondary",
                     use_container_width=True):
            st.session_state.page = pid
            st.rerun()
    st.divider()
    st.markdown("<div style='font-size: 11px; color: #64748b; font-weight: 500; margin-bottom: 8px;'>🟢 Offline (Ollama)</div>", unsafe_allow_html=True)
    if st.session_state.contract_name:
        if st.button("🗑 Clear & Start Over", use_container_width=True):
            for k, v in _defaults.items():
                st.session_state[k] = v
            try:
                from modules.embedder import clear_collection
                clear_collection()
            except ImportError:
                pass
            st.rerun()
    st.divider()
    st.caption("DATS 6501 Capstone | GWU")

page = st.session_state.page

# ══════════════════════════════════════════════════════════════════════════════
# UPLOAD
# ══════════════════════════════════════════════════════════════════════════════
if page == "upload":
    st.markdown("""
        <div style='text-align: center; padding: 3rem 0 2rem 0;'>
            <h1 style='font-size: 3.5rem; color: #3b82f6; margin-bottom: 0.5rem;'>AI Contract Review</h1>
            <p style='font-size: 1.25rem; color: #9ca3af; font-weight: 400;'>Instantly analyze, extract, and chat with your legal documents using advanced AI.</p>
        </div>
    """, unsafe_allow_html=True)
    st.divider()

    if not os.getenv("GROQ_API_KEY"):
        st.error("Groq API key not set. Add `GROQ_API_KEY=gsk_...` to your `.env` file and restart.")
        st.stop()

    uploaded_file = st.file_uploader(
        "Drag and drop your contract here, or click Browse",
        type=["pdf", "docx"], help="Supported: PDF, Word (.docx)"
    )

    if uploaded_file:
        st.markdown(f"**File:** {uploaded_file.name}  |  **Size:** {uploaded_file.size // 1024} KB")
        if st.button("Analyze Contract", type="primary"):
            import tempfile
            from modules.doc_processor import process_contract
            from modules.topic_modeler import classify_chunks_llm
            from modules.clause_matcher import match_chunks
            from modules.embedder import index_chunks
            from modules.llm_engine import extract_clauses

            suffix = ".pdf" if uploaded_file.name.endswith(".pdf") else ".docx"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name

            st.session_state.contract_name = uploaded_file.name
            st.session_state.chat_history = []  # Reset chat history for the new contract
            st.session_state.audit_log = []
            st.session_state.hitl_decisions = {}
            prog = st.progress(0)
            status = st.empty()

            try:
                status.info("📄 Step 1/6 — Reading document...")
                prog.progress(10)
                chunks = process_contract(tmp_path, chunk_size=500, overlap=80)
                # FIX: Override temp filepath with the actual filename for proper RAG filtering
                for chunk in chunks:
                    chunk.source_file = uploaded_file.name
                st.session_state.chunks = chunks
                chunk_texts = [c.text for c in chunks]

                status.info("🛡️ Step 2/6 — Verifying document type...")
                prog.progress(20)
                from modules.llm_engine import verify_document_type
                verification = verify_document_type(" ".join(chunk_texts[:3]))
                if not verification.get("is_contract", True):
                    st.error(f"🛑 **Document Rejected:** This doesn't appear to be a contract. {verification.get('reason')}")
                    status.empty()
                    prog.empty()
                    st.stop()


                status.info("🔗 Step 4/6 — Matching clauses to legal taxonomy...")
                prog.progress(55)
                match_results = match_chunks(chunk_texts)
                st.session_state.match_results = match_results

                status.info("💾 Step 5/6 — Building search index...")
                prog.progress(75)
                index_chunks(chunks, topic_labels=match_results)
                st.session_state.indexed = True

                status.info("🤖 Step 6/6 — Running full analysis...")
                prog.progress(90)
                analysis = extract_clauses(" ".join(chunk_texts[:8]), contract_name=uploaded_file.name)
                st.session_state.contract_analysis = analysis
                prog.progress(100)
                status.empty()

            except Exception as e:
                st.error(f"Something went wrong: {e}")
                st.stop()
            finally:
                os.unlink(tmp_path)

            # ── FIX 1: navigate directly, no extra button needed ─────────────
            st.session_state.page = "review"
            st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# REVIEW CONTRACT
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# REVIEW CONTRACT
# ══════════════════════════════════════════════════════════════════════════════
elif page == "review":
    st.title("🔍 Review Contract")

    if not st.session_state.contract_analysis:
        st.info("Please upload a contract first.")
        if st.button("Go to Upload"):
            st.session_state.page = "upload"
            st.rerun()
        st.stop()

    analysis  = st.session_state.contract_analysis
    decisions = st.session_state.hitl_decisions
    clauses   = analysis.clauses
    risk_emoji = {"low": "🟢", "medium": "🟡", "high": "🔴"}.get(analysis.overall_risk_level.lower(), "⚪")
    reviewed = sum(1 for c in clauses if c.clause_type in decisions)

    # Routing Logic
    routing_departments = set(["Legal"])
    for c in clauses:
        text = (c.clause_type + " " + getattr(c, "summary", "")).lower()
        if any(w in text for w in ["data", "privacy", "security", "confidential", "sharing"]):
            routing_departments.add("Security")
        if any(w in text for w in ["payment", "fee", "liability", "indemnity", "financial", "cap"]):
            routing_departments.add("Finance")
    
    routing_html = "".join([f"<span class='source-tag' style='font-size:13px; font-weight:600;'>{dept}</span>" for dept in sorted(list(routing_departments))])

    st.markdown(f"**Contract:** {analysis.contract_name}")
    c1, c2 = st.columns([0.4, 0.6])
    
    risk_class = f"risk-{analysis.overall_risk_level.lower()}"
    with c1:
        st.markdown(f"**Overall Risk Profile**")
        st.markdown(f"<div class='risk-pill {risk_class}' style='font-size: 1.2rem; padding: 8px 12px; text-align: center; width: 100%; background: rgba(255,255,255,0.05); border-radius: 8px; border: 1px solid rgba(255,255,255,0.1);'>{risk_emoji} {analysis.overall_risk_level.upper()}</div>", unsafe_allow_html=True)
    
    with c2:
        st.markdown(f"**Suggested Routing**")
        st.markdown(f"<div style='padding: 10px 0;'>{routing_html}</div>", unsafe_allow_html=True)

    st.divider()

    st.markdown(f"### 📄 Contract Summary")
    st.markdown(f"**TL;DR:** {getattr(analysis, 'executive_summary', 'No executive summary provided by model.')}")
    st.divider()

    # Add Red Flags and Missing Clauses to Review page for better visibility
    if analysis.red_flags:
        st.subheader("⚠️ Critical Red Flags")
        for f in analysis.red_flags:
            st.markdown(f"<div class='flag-card'>🚩 {f}</div>", unsafe_allow_html=True)
        st.divider()

    if analysis.missing_clauses:
        st.subheader("🔴 Missing Standard Clauses")
        for m in analysis.missing_clauses:
            st.markdown(f"<div class='missing-card'>⚠️ Not found: {m}</div>", unsafe_allow_html=True)
        st.divider()

    st.subheader("Key Clauses Extracted")
    st.caption("Review each clause. Edit if needed, then Accept or Reject.")

    # Use a cleaner one-column layout for review
    for i, clause in enumerate(clauses):
        icon  = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(clause.risk_level, "⚪")
        dec   = decisions.get(clause.clause_type, {})
        badge = {"accepted": " ✅", "edited": " ✏️", "rejected": " ❌"}.get(dec.get("decision", ""), "")
        with st.expander(
            f"{icon} **{clause.clause_type.title()}**{badge}  —  [Model Conf: {clause.confidence:.1%} | {len(clause.extracted_text.split())} tokens]",
            expanded=(clause.risk_level == "high" and clause.clause_type not in decisions)
        ):
            st.markdown(f"*{clause.summary}*")
            st.caption(f"Risk: **{clause.risk_level.upper()}** — {clause.risk_explanation}")
            edited = st.text_area("Clause text (you can edit)",
                                  value=clause.extracted_text, height=90,
                                  key=f"ct_{i}", label_visibility="collapsed")
            a1, a2, a3 = st.columns(3)
            with a1:
                if st.button("✅ Accept Validation", key=f"acc_{i}", use_container_width=True):
                    decisions[clause.clause_type] = {"decision": "accepted", "edited_text": edited}
                    st.session_state.audit_log.append({
                        "action": f"**{clause.clause_type}** accepted",
                        "time": datetime.datetime.now().strftime("%I:%M %p")})
                    st.rerun()
            with a2:
                if st.button("✏️ Edit & Save", key=f"edit_{i}", use_container_width=True):
                    decisions[clause.clause_type] = {"decision": "edited", "edited_text": edited}
                    st.session_state.audit_log.append({
                        "action": f"**{clause.clause_type}** edited",
                        "time": datetime.datetime.now().strftime("%I:%M %p")})
                    st.rerun()
            with a3:
                if st.button("❌ False Positive", key=f"rej_{i}", use_container_width=True):
                    decisions[clause.clause_type] = {"decision": "rejected", "edited_text": ""}
                    st.session_state.audit_log.append({
                        "action": f"**{clause.clause_type}** rejected",
                        "time": datetime.datetime.now().strftime("%I:%M %p")})
                    st.rerun()
            if clause.clause_type in decisions:
                st.caption({"accepted": "✅ Accepted", "edited": "✏️ Edited & Accepted",
                            "rejected": "❌ Rejected"}.get(dec.get("decision", ""), ""))

    st.divider()
    st.subheader("📥 Export Final Report")
    st.caption("Generate a downloadable DOCX report containing the extracted clauses, risk analysis, and review decisions.")
    
    def create_docx_report():
        import io
        from docx import Document
        from docx.shared import RGBColor
        
        doc = Document()
        doc.add_heading(f"Contract Review Report: {analysis.contract_name}", 0)
        
        doc.add_heading("Overall Risk Profile", level=1)
        doc.add_paragraph(f"Risk Level: {analysis.overall_risk_level.upper()}")
        doc.add_paragraph(analysis.executive_summary)

        # Routing tags
        routine_tags = ", ".join(list(routing_departments)) if routing_departments else "Legal"
        doc.add_heading("Suggested Routing", level=1)
        doc.add_paragraph(routine_tags)
        
        doc.add_heading("Extracted Clauses", level=1)
        for clause in clauses:
            dec = decisions.get(clause.clause_type, {})
            status = dec.get("decision", "pending")
            if status == "rejected":
                continue
                
            p = doc.add_paragraph()
            p.add_run(f"{clause.clause_type.title()} ").bold = True
            run = p.add_run(f"({status.upper()})")
            if clause.risk_level == 'high':
                run.font.color.rgb = RGBColor(255, 0, 0)
            
            doc.add_paragraph(f"Risk Level: {clause.risk_level.upper()}")
            doc.add_paragraph(f"Summary: {clause.summary}")
            
            text_to_show = dec.get("edited_text", clause.extracted_text)
            if text_to_show:
                doc.add_paragraph(f"Extracted Text:\n{text_to_show}")
            
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    st.download_button(
        label="📄 Generate & Download DOCX Report",
        data=create_docx_report(),
        file_name=f"ReviewReport_{analysis.contract_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )

# ══════════════════════════════════════════════════════════════════════════════
# Q&A PAGE
# ══════════════════════════════════════════════════════════════════════════════
elif page == "qa":
    st.title("💬 Contract Q&A")

    if not st.session_state.contract_analysis:
        st.info("Please upload a contract first.")
        if st.button("Go to Upload"):
            st.session_state.page = "upload"
            st.rerun()
        st.stop()

    st.markdown(f"Ask anything about **{st.session_state.contract_name}** in plain English.")
    st.divider()

    # Initialize dynamic questions state if absent
    if "dynamic_questions" not in st.session_state:
        st.session_state.dynamic_questions = []

    # If dynamic questions haven't been loaded, generate them
    if not st.session_state.dynamic_questions and st.session_state.contract_analysis:
        from modules.llm_engine import generate_dynamic_questions
        with st.spinner("Generating custom contract queries..."):
            # Provide REAL text from the contract to get highly specific questions
            if "chunks" in st.session_state and st.session_state.chunks:
                source_text = "\n".join([c.text for c in st.session_state.chunks[:4]])
            else:
                source_text = st.session_state.contract_analysis.executive_summary
                
            st.session_state.dynamic_questions = generate_dynamic_questions(source_text)
            if not st.session_state.dynamic_questions:
                # fallback
                st.session_state.dynamic_questions = [
                    "What is the liability cap?",
                    "How can this contract be terminated?",
                    "Who owns the intellectual property?"
                ]

    # Suggested questions — more compact design
    st.markdown("<div style='font-size: 0.85rem; font-weight: 500; color: #60a5fa;'>SUGGESTED QUERIES</div>", unsafe_allow_html=True)
    suggestions = st.session_state.dynamic_questions
    
    cols = st.columns(3)
    for idx, full_q in enumerate(suggestions[:3]):
        # Display the complete question without truncation
        if cols[idx % 3].button(full_q, key=f"sq_{idx}", use_container_width=True):
            st.session_state.prefill_q = full_q
            st.session_state.auto_ask  = True

    st.divider()

    # Voice section
    col_v, _ = st.columns([1, 2])
    with col_v:
        st.markdown("**🎤 Voice - click mic and speak**")
        try:
            from streamlit_mic_recorder import mic_recorder
            audio_in = mic_recorder(start_prompt="🎤 Start recording",
                                    stop_prompt="⏹ Stop", key="mic_main", just_once=True)
            if audio_in and audio_in.get("bytes"):
                from modules.voice_handler import transcribe_audio_bytes, is_valid_question
                with st.spinner("Transcribing..."):
                    try:
                        spoken, _ = transcribe_audio_bytes(audio_in["bytes"], audio_format="wav")
                        if is_valid_question(spoken):
                            st.session_state.prefill_q = spoken
                            st.session_state.auto_ask  = True
                            st.success(f'Heard: "{spoken}"')
                            st.rerun()
                        else:
                            st.warning("Didn't catch that, try again.")
                    except Exception as e:
                        st.error(f"Voice error: {e}")
        except ImportError:
            st.caption("Voice not installed.")

    st.divider()

    # Initialize edit state
    if "editing_msg_idx" not in st.session_state:
        st.session_state.editing_msg_idx = None

    # Render Chat History
    for idx, msg in enumerate(st.session_state.chat_history):
        if msg["role"] == "user":
            if st.session_state.editing_msg_idx == idx:
                with st.form(key=f"edit_form_{idx}"):
                    edit_q = st.text_input("Edit question", value=msg["content"], label_visibility="collapsed")
                    ec1, ec2 = st.columns([1,1])
                    if ec1.form_submit_button("✅ Save"):
                        st.session_state.chat_history = st.session_state.chat_history[:idx]
                        st.session_state.prefill_q = edit_q
                        st.session_state.auto_ask = True
                        st.session_state.editing_msg_idx = None
                        st.rerun()
                    if ec2.form_submit_button("❌"):
                        st.session_state.editing_msg_idx = None
                        st.rerun()
            else:
                st.markdown(f"""
                <div class='user-msg'>
                    <div style='display: flex; justify-content: space-between;'>
                        <div class='msg-label'>YOU</div>
                    </div>
                    <div>{msg['content']}</div>
                </div>
                """, unsafe_allow_html=True)
        else:
            # Assistant Message with premium styling
            sources_html = "".join([f"<span class='source-tag'>{s}</span>" for s in msg.get("sources", [])[:3]])
            conf_color = "#10b981" if msg.get("confidence", 0) > 0.7 else "#f59e0b" if msg.get("confidence", 0) > 0.4 else "#ef4444"
            
            st.markdown(f"""
            <div class='bot-msg'>
                <div class='msg-label'>AI ASSISTANT</div>
                <div style='margin-bottom: 12px;'>{msg['content'].replace('**', '')}</div>
                <div style='display: flex; align-items: center; gap: 10px; opacity: 0.8;'>
                    <div style='font-size: 11px; color: {conf_color}; font-weight: 700;'>
                        CONFIDENCE: {msg.get('confidence', 0):.0%}
                    </div>
                    <div style='border-left: 1px solid rgba(255,255,255,0.1); height: 12px;'></div>
                    <div>{sources_html}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            if msg.get("audio_out"):
                st.audio(msg["audio_out"])
            
            if msg.get("chunks_used"):
                with st.expander("🔍 Legal Citations", expanded=False):
                    for i, chunk in enumerate(msg["chunks_used"][:3]):
                        st.caption(f"Context Block {i+1}")
                        st.text(chunk["text"][:300] + "...")
                        if i < 2: st.divider()

    if st.session_state.chat_history:
        if st.button("🗑️", help="Clear Conversation"):
            st.session_state.chat_history = []
            st.rerun()

    # Auto-scroll to bottom and button placement
    import streamlit.components.v1 as components
    components.html(
        """
        <script>
            const doc = window.parent.document;
            
            // Auto-scroll logic
            const scrollNode = doc.querySelector('.main') || doc.querySelector('.stApp');
            if (scrollNode) {
                window.parent.requestAnimationFrame(() => {
                    scrollNode.scrollTo({top: scrollNode.scrollHeight, behavior: 'smooth'});
                });
            }

            // Position clear button inside the chat input box
            function enhanceChat() {
                const buttons = doc.querySelectorAll('button');
                let clearBtn = null;
                buttons.forEach(btn => {
                    if (btn.innerText.includes('🗑️')) {
                        clearBtn = btn;
                    }
                });
                
                const chatInput = doc.querySelector('div[data-testid="stChatInput"]');
                if (clearBtn && chatInput) {
                    const chatRect = chatInput.getBoundingClientRect();
                    
                    // Hide the wrapper so it doesn't take up space in the normal layout
                    const wrapper = clearBtn.closest('div[data-testid="element-container"]');
                    if (wrapper && wrapper.style.position !== 'absolute') {
                        wrapper.style.position = 'absolute';
                        wrapper.style.visibility = 'hidden';
                    }
                    
                    // Detach button visually
                    clearBtn.style.position = 'fixed';
                    clearBtn.style.visibility = 'visible';
                    clearBtn.style.zIndex = '99999';
                    
                    // Position it inside the right side of the chat box, just left of the send button
                    clearBtn.style.left = (chatRect.right - 90) + 'px';
                    clearBtn.style.top = (chatRect.top + chatRect.height / 2 - 18) + 'px';
                    
                    // Premium styling
                    clearBtn.style.borderRadius = '50%';
                    clearBtn.style.width = '36px';
                    clearBtn.style.height = '36px';
                    clearBtn.style.background = 'rgba(30, 41, 59, 0.9)';
                    clearBtn.style.border = '1px solid rgba(255, 255, 255, 0.1)';
                    clearBtn.style.display = 'flex';
                    clearBtn.style.alignItems = 'center';
                    clearBtn.style.justifyContent = 'center';
                    clearBtn.style.boxShadow = '0 2px 8px rgba(0,0,0,0.3)';
                    clearBtn.style.cursor = 'pointer';
                    
                    // Add hover effects using JS since we're manipulating styles directly
                    clearBtn.onmouseover = () => {
                        clearBtn.style.background = 'rgba(239, 68, 68, 0.2)';
                        clearBtn.style.borderColor = 'rgba(239, 68, 68, 0.5)';
                    };
                    clearBtn.onmouseout = () => {
                        clearBtn.style.background = 'rgba(30, 41, 59, 0.9)';
                        clearBtn.style.borderColor = 'rgba(255, 255, 255, 0.1)';
                    };
                    
                    // Pad the textarea so text doesn't hide behind our button
                    const textArea = chatInput.querySelector('textarea');
                    if (textArea) {
                        textArea.style.paddingRight = '100px';
                    }
                }
            }
            
            enhanceChat();
            window.parent.addEventListener('resize', enhanceChat);
            setInterval(enhanceChat, 500);
        </script>
        """,
        height=0
    )

    # The chat input
    prompt = st.chat_input("Ask a question about the contract...")
    should_answer = prompt or (st.session_state.auto_ask and st.session_state.prefill_q)

    if should_answer:
        final_question = prompt if prompt else st.session_state.prefill_q
        st.session_state.prefill_q = ""
        st.session_state.auto_ask  = False
        
        st.session_state.chat_history.append({"role": "user", "content": final_question})
        
        from modules.rag_pipeline import rag_retrieve_and_assemble
        from modules.llm_engine import answer_question

        with st.spinner("Analyzing contract text..."):
            try:
                context, chunks_used = rag_retrieve_and_assemble(
                    query=final_question, 
                    top_k=3,
                    filter_source_file=st.session_state.contract_name)
                answer = answer_question(
                    question=final_question,
                    context=context,
                    chat_history=st.session_state.chat_history,
                    chunks_used=chunks_used)
                
                audio_out = None
                try:
                    from modules.voice_handler import text_to_speech, answer_for_voice
                    audio_out = text_to_speech(answer_for_voice(answer))
                except Exception: pass

                st.session_state.chat_history.append({
                    "role": "assistant", 
                    "content": answer.answer,
                    "confidence": answer.confidence,
                    "sources": answer.sources,
                    "audio_out": audio_out,
                    "chunks_used": chunks_used
                })
                st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

    
    # ══════════════════════════════════════════════════════════════════════════════
    # DASHBOARD — FIX 4: imports inside block, safe .get() everywhere
    # ══════════════════════════════════════════════════════════════════════════════
elif page == "dashboard":
    st.markdown("<h1 class='executive-header'>Risk Intelligence Dashboard</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color: #94a3b8; margin-bottom: 30px;'>Multi-dimensional audit analytics and vector-based risk assessment.</p>", unsafe_allow_html=True)

    if not st.session_state.match_results:
        st.info("Upload a contract to see the dashboard.")
        if st.button("Go to Upload"):
            st.session_state.page = "upload"
            st.rerun()
        st.stop()

    import plotly.express as px
    import plotly.graph_objects as go
    import pandas as pd
    from modules.clause_matcher import summarize_risk_profile

    mr       = st.session_state.match_results
    profile  = summarize_risk_profile(mr)
    analysis = st.session_state.contract_analysis
    re       = {"low": "🟢", "medium": "🟡", "high": "🔴"}.get(
        analysis.overall_risk_level.lower(), "⚪") if analysis else "⚪"

    # 1. PREMIUM KPI ROW
    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.markdown(f"""
        <div class='glass-card' style='padding: 20px; text-align: center; height: 135px; display: flex; flex-direction: column; justify-content: center;'>
            <div class='metric-label-small' style='font-size: 0.7rem; letter-spacing: 1.5px; margin-bottom: 8px;'>OVERALL PROFILE</div>
            <div class='metric-value-large' style='color: {"#f43f5e" if analysis.overall_risk_level=="high" else "#10b981"}; margin-bottom: 6px;'>
                {analysis.overall_risk_level[0].upper()}
            </div>
            <div style='font-size: 0.75rem; font-weight: 600; color: #64748b; letter-spacing: 0.5px;'>{analysis.overall_risk_level.upper()} RISK</div>
        </div>
        """, unsafe_allow_html=True)
    with k2:
        st.markdown(f"""
        <div class='glass-card' style='padding: 20px; text-align: center; height: 135px; display: flex; flex-direction: column; justify-content: center;'>
            <div class='metric-label-small' style='font-size: 0.7rem; letter-spacing: 1.5px; margin-bottom: 8px;'>SECTIONS SCANNED</div>
            <div class='metric-value-large' style='margin-bottom: 6px;'>{profile.get("total_clauses", 0)}</div>
            <div style='font-size: 0.75rem; font-weight: 600; color: #64748b; letter-spacing: 0.5px;'>VECTOR CHUNKS</div>
        </div>
        """, unsafe_allow_html=True)
    with k3:
        st.markdown(f"""
        <div class='glass-card' style='padding: 20px; text-align: center; height: 135px; display: flex; flex-direction: column; justify-content: center;'>
            <div class='metric-label-small' style='font-size: 0.7rem; letter-spacing: 1.5px; margin-bottom: 8px;'>AGGREGATE SCORE</div>
            <div class='metric-value-large' style='margin-bottom: 6px;'>{profile.get('average_risk_score', 0):.2f}</div>
            <div style='font-size: 0.75rem; font-weight: 600; color: #64748b; letter-spacing: 0.5px;'>MAX 1.00</div>
        </div>
        """, unsafe_allow_html=True)
    with k4:
        clauses = st.session_state.contract_analysis.clauses if st.session_state.contract_analysis else []
        decisions = st.session_state.get("hitl_decisions", {})
        reviewed = sum(1 for c in clauses if c.clause_type in decisions)
        rate_p = reviewed / max(len(clauses), 1)
        st.markdown(f"""
        <div class='glass-card' style='padding: 20px; text-align: center; height: 135px; display: flex; flex-direction: column; justify-content: center;'>
            <div class='metric-label-small' style='font-size: 0.7rem; letter-spacing: 1.5px; margin-bottom: 8px;'>VERIFICATION</div>
            <div class='metric-value-large' style='margin-bottom: 6px;'>{rate_p:.0%}</div>
            <div style='font-size: 0.75rem; font-weight: 600; color: #64748b; letter-spacing: 0.5px;'>AUDIT PROGRESS</div>
        </div>
        """, unsafe_allow_html=True)

    # 2. ANALYST SUMMARY CARD
    if analysis:
        st.markdown(f"""
        <div class='glass-card' style='padding: 24px 30px; border-left: 3px solid #3b82f6; background: linear-gradient(135deg, rgba(15, 23, 42, 0.6) 0%, rgba(30, 41, 59, 0.4) 100%); border-radius: 16px;'>
            <div class='msg-label' style='color: #60a5fa; margin-bottom: 12px; font-size: 10px; letter-spacing: 2px;'>PRIMARY RISK ASSESSMENT</div>
            <div style='font-size: 1.05rem; line-height: 1.7; color: #cbd5e1; font-weight: 300;'>
                {analysis.executive_summary}
            </div>
        </div>
        """, unsafe_allow_html=True)
    # 3. INTUITIVE ANALYTICS
    st.markdown("<br>", unsafe_allow_html=True)
    v1 = st.container()
    import streamlit.components.v1 as components

    with v1:
        c0, c2 = st.columns([0.8, 0.2])
        with c0:
            st.markdown("<div class='msg-label' style='margin-bottom: 20px; color: #60a5fa;'>CONTRACT RISK METER</div>", unsafe_allow_html=True)
        tech_mode = False  # Disabled UI toggle to enforce standard business metrics
        with c2:
            components.html("""
            <style>
                .pdf-btn {
                    padding: 4px 10px;
                    background: linear-gradient(135deg, rgba(59, 130, 246, 0.2), rgba(30, 64, 175, 0.2));
                    border: 1px solid rgba(96, 165, 250, 0.5);
                    color: #93c5fd;
                    border-radius: 6px;
                    font-family: -apple-system, system-ui, sans-serif;
                    font-size: 13px;
                    font-weight: 500;
                    cursor: pointer;
                    width: 100%;
                    transition: all 0.2s ease;
                }
                .pdf-btn:hover {
                    background: rgba(59, 130, 246, 0.4);
                    color: #ffffff;
                }
            </style>
            <button class="pdf-btn" onclick="window.parent.print()">
                📄 PDF
            </button>
            """, height=40)

        if mr:
            import numpy as np
            # 1. Calculate Risk Score
            avg_risk = profile.get('average_risk_score', 0)
            llm_risk_level = getattr(analysis, 'overall_risk_level', 'medium').lower()
            
            # Anchor Risk Score strictly to LLM holistic assessment
            if llm_risk_level == "high":
                risk_score = int(min(100, max(75, 75 + (avg_risk * 25))))
                risk_str = "High"
            elif llm_risk_level == "low":
                risk_score = int(min(30, max(0, avg_risk * 30)))
                risk_str = "Low"
            else:
                risk_score = int(min(74, max(31, 31 + (avg_risk * 38))))
                risk_str = "Medium"
            
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=risk_score,
                number={'suffix': "/100", 'font': {'size': 48, 'color': '#f8fafc', 'family': 'Outfit'}},
                domain={'x': [0, 1], 'y': [0, 1]},
                gauge={
                    'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "rgba(255,255,255,0.2)"},
                    'bar': {'color': "rgba(255,255,255,0.9)", 'thickness': 0.15},
                    'bgcolor': "rgba(255,255,255,0.05)",
                    'borderwidth': 0,
                    'steps': [
                        {'range': [0, 30], 'color': "rgba(16, 185, 129, 0.75)"},   # Green
                        {'range': [30, 75], 'color': "rgba(251, 191, 36, 0.75)"},  # Yellow
                        {'range': [75, 100], 'color': "rgba(244, 63, 94, 0.75)"}   # Red
                    ],
                    'threshold': {
                        'line': {'color': "#ffffff", 'width': 4},
                        'thickness': 0.75,
                        'value': risk_score
                    }
                }
            ))
            
            fig.update_layout(
                height=220, margin=dict(l=20, r=20, t=20, b=20),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                font=dict(family="Outfit", color="#94a3b8")
            )
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})
            
            # --- 2. Additional Breakdown Cards (Two-Track Logic) ---
            df_full = pd.DataFrame(mr)
            found_count = len(df_full['matched_clause'].unique())
            completeness = min(100, int((found_count / 14) * 100)) # 14 canonical clauses standard
            
            try:
                avg_chunk_len = int(np.mean([len(str(c.get('text', ''))) for c in mr]))
                clarity = max(40, min(98, 100 - int(avg_chunk_len / 50)))
            except:
                clarity = 78
                
            risk_col = "#fca5a5" if risk_str=="High" else ("#fde047" if risk_str=="Medium" else "#6ee7b7")

            # Dynamic Dictionaries for display modes
            if completeness >= 90:
                b_comp_t, b_comp_v = "✅ COMPLETENESS", "All standard clauses included ✓"
            elif completeness >= 60:
                b_comp_t, b_comp_v = "✅ COMPLETENESS", f"Coverage: {completeness}%"
            else:
                b_comp_t, b_comp_v = "❌ COMPLETENESS", "Missing key clauses"
            b_comp_s = getattr(analysis, "completeness_assessment", "Dynamic completeness overview (Re-upload contract to generate).")
            t_comp_t, t_comp_v, t_comp_s = "✅ COMPLETENESS (NLP)", f"{completeness}% Match", f"Found: {found_count} / 14 canonical vectors. Missing: {len(analysis.missing_clauses or [])} core components."

            if clarity >= 80:
                b_clar_t, b_clar_v = "💎 READABILITY", "Highly Readable"
            elif clarity >= 60:
                b_clar_t, b_clar_v = "💎 READABILITY", "Moderate"
            else:
                b_clar_t, b_clar_v = "💎 READABILITY", "Complex"
            b_clar_s = getattr(analysis, "readability_assessment", "Dynamic linguistic overview (Re-upload contract to generate).")
            t_clar_t, t_clar_v, t_clar_s = "🔍 CLARITY SCORE", f"{clarity}%", f"Avg semantic cluster span: {avg_chunk_len} chars."

            if risk_str == "Low":
                b_risk_t, b_risk_v = "⚠️ RISK LEVEL", "Safe (Low Risk)"
            elif risk_str == "Medium":
                b_risk_t, b_risk_v = "⚠️ RISK LEVEL", "Moderate Risk"
            else:
                b_risk_t, b_risk_v = "⚠️ RISK LEVEL", "High Risk"
            b_risk_s = getattr(analysis, "risk_assessment", "Dynamic risk overview (Re-upload contract to generate).")
            t_risk_t, t_risk_v, t_risk_s = "⚠️ RISK LEVEL", f"{risk_str} ({avg_risk:.2f})", f"Aggregate ML risk mapped across {found_count} logic vectors."

            if risk_score > 70:
                b_bal_t, b_bal_v = "⚖️ FAIRNESS", "One-Sided"
            elif risk_score <= 40:
                b_bal_t, b_bal_v = "⚖️ FAIRNESS", "Balanced"
            else:
                b_bal_t, b_bal_v = "⚖️ FAIRNESS", "Slightly Skewed"
            b_bal_s = getattr(analysis, "fairness_assessment", "Dynamic fairness overview (Re-upload contract to generate).")
            balance_tech = "Company Skew" if risk_score > 70 else ("Neutral" if risk_score <= 40 else "Warning")
            t_bal_t, t_bal_v, t_bal_s = "⚖️ BALANCE BIAS", balance_tech, "Analyzed via asymmetry in risk matrices."

            # Render logic
            comp_t, comp_v, comp_s = (t_comp_t, t_comp_v, t_comp_s) if tech_mode else (b_comp_t, b_comp_v, b_comp_s)
            clar_t, clar_v, clar_s = (t_clar_t, t_clar_v, t_clar_s) if tech_mode else (b_clar_t, b_clar_v, b_clar_s)
            risk_t, risk_v, risk_s = (t_risk_t, t_risk_v, t_risk_s) if tech_mode else (b_risk_t, b_risk_v, b_risk_s)
            bal_t, bal_v, bal_s = (t_bal_t, t_bal_v, t_bal_s) if tech_mode else (b_bal_t, b_bal_v, b_bal_s)

            value_font_size = "18px" if not tech_mode else "20px"
            info_icon = "<span style='display:inline-block; border: 1px solid rgba(148,163,184,0.5); border-radius: 50%; width: 14px; height: 14px; text-align: center; line-height: 13px; font-size: 9px; font-weight: 600; font-family: sans-serif; margin-left: 4px; cursor: help;'>i</span>"

            m1, m2 = st.columns(2)
            with m1:
                st.markdown(f"""
                <div title='Percentage of standard contract clauses successfully identified.' style='background: rgba(15, 23, 42, 0.5); border: 1px solid rgba(255,255,255,0.05); padding: 15px 20px; border-radius: 12px; margin-bottom: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.2);'>
                    <div style='font-size: 11px; color: #94a3b8; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 4px; display: flex; align-items: center;'>{comp_t}{info_icon}</div>
                    <div style='font-size: {value_font_size}; color: #f8fafc; font-weight: 700;'>{comp_v}</div>
                    <div style='font-size: 12px; color: #94a3b8; margin-top: 4px; line-height: 1.4;'>{comp_s}</div>
                </div>
                """, unsafe_allow_html=True)
                st.markdown(f"""
                <div title='Readability metric derived from average semantic clause span (legalese density).' style='background: rgba(15, 23, 42, 0.5); border: 1px solid rgba(255,255,255,0.05); padding: 15px 20px; border-radius: 12px; margin-bottom: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.2);'>
                    <div style='font-size: 11px; color: #94a3b8; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 4px; display: flex; align-items: center;'>{clar_t}{info_icon}</div>
                    <div style='font-size: {value_font_size}; color: #f8fafc; font-weight: 700;'>{clar_v}</div>
                    <div style='font-size: 12px; color: #94a3b8; margin-top: 4px; line-height: 1.4;'>{clar_s}</div>
                </div>
                """, unsafe_allow_html=True)
            with m2:
                st.markdown(f"""
                <div title='Calculated aggregate risk severity indicating signing safety.' style='background: rgba(15, 23, 42, 0.5); border: 1px solid rgba(255,255,255,0.05); border-right: 3px solid {risk_col}; padding: 15px 20px; border-radius: 12px; margin-bottom: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.2);'>
                    <div style='font-size: 11px; color: #94a3b8; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 4px; display: flex; align-items: center;'>{risk_t}{info_icon}</div>
                    <div style='font-size: {value_font_size}; color: {risk_col}; font-weight: 700;'>{risk_v}</div>
                    <div style='font-size: 12px; color: #94a3b8; margin-top: 4px; line-height: 1.4;'>{risk_s}</div>
                </div>
                """, unsafe_allow_html=True)
                st.markdown(f"""
                <div title='Asymmetry analysis determining if the contract heavily favors the counterparty.' style='background: rgba(15, 23, 42, 0.5); border: 1px solid rgba(255,255,255,0.05); padding: 15px 20px; border-radius: 12px; margin-bottom: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.2);'>
                    <div style='font-size: 11px; color: #94a3b8; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 4px; display: flex; align-items: center;'>{bal_t}{info_icon}</div>
                    <div style='font-size: {value_font_size}; color: #f8fafc; font-weight: 700;'>{bal_v}</div>
                    <div style='font-size: 12px; color: #94a3b8; margin-top: 4px; line-height: 1.4;'>{bal_s}</div>
                </div>
                """, unsafe_allow_html=True)



    # 4. CAPSTONE-WORTHY RED FLAGS & VULNERABILITIES
    if analysis.red_flags or analysis.missing_clauses:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<div class='msg-label' style='color: #f43f5e; margin-bottom: 5px; font-size: 14px;'>IDENTIFIED LEGAL VULNERABILITIES</div>", unsafe_allow_html=True)
        st.markdown("<div style='height: 2px; width: 100%; background: linear-gradient(90deg, #f43f5e 0%, transparent 100%); margin-bottom: 20px;'></div>", unsafe_allow_html=True)
        
        for f in analysis.red_flags:
            fw = f.lower()
            if any(w in fw for w in ["missing", "failure", "breach", "liability", "indemnif", "cap"]):
                icon, level, color = "🔴", "CRITICAL", "#fca5a5"
                bg_col = "rgba(252, 165, 165, 0.05)"
            elif any(w in fw for w in ["unclear", "ambiguous", "notice", "term", "auto-renewal"]):
                icon, level, color = "🟡", "WARNING", "#fde047"
                bg_col = "rgba(253, 224, 71, 0.05)"
            else:
                icon, level, color = "🟠", "MODERATE", "#fdba74"
                bg_col = "rgba(253, 186, 116, 0.05)"
                
            # 2. Heuristic Recommendation Mapping
            rec = "Ensure manual review by legal counsel prior to signature."
            if level == "CRITICAL":
                rec = "Stop execution. Flag immediately to the executive team for renegotiation of core terms."
            elif any(w in fw for w in ["financial", "loan", "payment", "fee", "cost"]):
                rec = "Have the Finance team review the payment schedule and default triggers."
            elif any(w in fw for w in ["termination", "default", "breach"]):
                rec = "Negotiate a 30-day cure period (grace period) for non-payment defaults."
            elif any(w in fw for w in ["unclear", "ambiguous", "vague"]):
                rec = "Mandate Legal to explicitly define these vague technical terms to prevent litigation."

            # 3. Best Clause Match (Simulated Context)
            clause_match = None
            for c in clauses:
                if c.risk_level.lower() == level.lower() or ("high" in c.risk_level.lower() and level=="CRITICAL"):
                    clause_match = c
                    break
            if not clause_match and clauses:
                clause_match = clauses[0]
                
            with st.container():
                st.markdown(f"""
                <div style='background: {bg_col}; border-left: 3px solid {color}; padding: 15px 20px; border-radius: 8px; margin-bottom: 5px; border: 1px solid rgba(255,255,255,0.05);'>
                    <div style='font-size: 1.05rem; color: #e2e8f0; line-height: 1.5;'>
                        {icon} <span style='color: {color}; font-weight: 700; font-size: 0.75rem; text-transform: uppercase; letter-spacing: 1px; margin-right: 8px;'>{level}</span> {f}
                    </div>
                </div>
                """, unsafe_allow_html=True)
                with st.expander("↳ View clause context"):
                    if tech_mode:
                        st.caption(f"Relevant context vector retrieved regarding: *{f}*. Review the extracted clauses tab or execute a semantic search query for exact phrasing.")
                    else:
                        st.markdown(f"**📍 Identified in Clause:** {clause_match.clause_type.title() if clause_match else 'General Provisions'}")
                        if clause_match:
                            st.info(f'"{clause_match.extracted_text[:180]}..."')
                        st.markdown(f"**⚠️ Actionable Recommendation:** {rec}")
        
        for m in (analysis.missing_clauses or []):
            with st.container():
                st.markdown(f"""
                <div style='background: rgba(252, 165, 165, 0.05); border-left: 3px solid #fca5a5; padding: 15px 20px; border-radius: 8px; margin-bottom: 5px; border: 1px solid rgba(255,255,255,0.05);'>
                    <div style='font-size: 1.05rem; color: #e2e8f0; line-height: 1.5;'>
                        🔴 <span style='color: #fca5a5; font-weight: 700; font-size: 0.75rem; text-transform: uppercase; letter-spacing: 1px; margin-right: 8px;'>CRITICAL</span> Missing explicitly standardized clause: <b>{m}</b>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                with st.expander("↳ Remediation Step"):
                    st.caption(f"This contract lacks a detected **{m}** clause. Consider routing to Legal for insertion.")

    st.markdown("<div style='margin-bottom: 100px;'></div>", unsafe_allow_html=True)
